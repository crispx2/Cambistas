from flask import Flask, render_template, request, jsonify, send_file
from docx import Document
from io import BytesIO

app = Flask(__name__)

data_store = []

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/adicionar', methods=['POST'])
def adicionar():
    data = request.json
    data_store.append(data)
    return jsonify({'status': 'ok'})

@app.route('/dados')
def dados():
    return jsonify(data_store)

@app.route('/apagar/<int:index>', methods=['DELETE'])
def apagar(index):
    try:
        data_store.pop(index)
        return jsonify({'status': 'ok'})
    except IndexError:
        return jsonify({'status': 'error', 'message': 'Índice inválido'}), 400

@app.route('/exportar-word', methods=['POST'])
def exportar_word():
    doc = Document()
    doc.add_heading('Relatório de Cambistas por Área', 0)

    areas = ['Palmeira dos Índio', 'Maragogi', 'Tamandaré']

    for area in areas:
        doc.add_heading(area, level=1)
        tabela_area = [d for d in data_store if d['area'] == area]

        if not tabela_area:
            doc.add_paragraph('Nenhum registro.')
            continue

        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Código'
        hdr_cells[1].text = 'Cambista'
        hdr_cells[2].text = 'Comissão (R$)'
        hdr_cells[3].text = 'Complemento de Dezena (R$)'

        for item in tabela_area:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get('codigo', ''))
            row_cells[1].text = str(item.get('cambista', ''))
            row_cells[2].text = f"R$ {float(item.get('comissao', 0)):,.2f}"
            row_cells[3].text = f"R$ {float(item.get('Complemento_De_Dezena', 0)):,.2f}"

        doc.add_paragraph('')  # espaço entre áreas

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='Relatorio_Cambistas.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

if __name__ == '__main__':
    app.run(debug=True)
